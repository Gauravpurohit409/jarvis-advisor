"""
Document Parser Service
Extracts client data from Word (.docx) and Excel (.xlsx) files
"""

import re
from datetime import date, datetime
from typing import Dict, Any, Optional, List, Tuple
from io import BytesIO


class DocumentParser:
    """
    Parser for extracting client data from Word and Excel documents.
    Uses pattern matching to find common client information fields.
    """
    
    # Required fields for a valid client
    REQUIRED_FIELDS = [
        "id", "title", "first_name", "last_name", "date_of_birth",
        "email", "phone", "address_line1", "city", "postcode"
    ]
    
    # Common field patterns to search for
    FIELD_PATTERNS = {
        "title": r"(?:title|salutation|mr/mrs)[:\s]*([A-Za-z]+)",
        "first_name": r"(?:first\s*name|forename|given\s*name)[:\s]*([A-Za-z\-]+)",
        "last_name": r"(?:last\s*name|surname|family\s*name)[:\s]*([A-Za-z\-]+)",
        "full_name": r"(?:full\s*name|name|client\s*name)[:\s]*([A-Za-z\-\s]+)",
        "date_of_birth": r"(?:date\s*of\s*birth|dob|birth\s*date|d\.o\.b)[:\s]*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})",
        "email": r"(?:email|e-mail)?[:\s]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})",
        "phone": r"(?:phone|telephone|tel|mobile|contact)?[:\s]*((?:07\d{3}|\+44\s*7\d{3})\s*\d{3}\s*\d{3,4})",
        "address": r"(?:address|street|residence)[:\s]*(.+?)(?=\n|$)",
        "city": r"(?:city|town)[:\s]*([A-Za-z\s\-]+)",
        "postcode": r"(?:postcode|post\s*code|zip)?[:\s]*([A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2})",
        "occupation": r"(?:occupation|job|profession)[:\s]+([A-Za-z][A-Za-z\s\-]{3,}?)(?=\n|$|,)",
        "employer": r"(?:employer|company|organisation|organization)[:\s]*([A-Za-z\s\-\&]+)",
        "income": r"[£$]([\d,]+)(?:\s*(?:per\s*year|p\.?a\.?|annually|/year)?)",
        "portfolio": r"(?:portfolio|net\s*worth|total\s*value|assets)[:\s]*[£$]?([\d,\.]+)",
        "marital": r"(?:marital|marriage|relationship)[:\s]*(single|married|divorced|widowed|civil)",
        "ni_number": r"(?:national\s*insurance|ni\s*number|nino)?[:\s]*([A-Z]{2}\d{6}[A-Z])",
    }
    
    # UK postcode pattern
    UK_POSTCODE = r"[A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2}"
    
    # Common titles
    VALID_TITLES = ["Mr", "Mrs", "Ms", "Miss", "Dr", "Prof", "Sir", "Dame"]
    
    def __init__(self):
        self._docx_available = False
        self._openpyxl_available = False
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if document parsing libraries are available"""
        try:
            import docx
            self._docx_available = True
        except ImportError:
            pass
        
        try:
            import openpyxl
            self._openpyxl_available = True
        except ImportError:
            pass
    
    def parse_document(self, file_content: bytes, filename: str) -> Tuple[Dict[str, Any], List[str]]:
        """
        Parse a document and extract client data.
        
        Args:
            file_content: Raw bytes of the uploaded file
            filename: Original filename to determine type
            
        Returns:
            Tuple of (extracted_data dict, list of missing required fields)
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.docx'):
            text = self._get_docx_text(file_content)
            extracted = self._extract_from_text(text)
        elif filename_lower.endswith('.xlsx') or filename_lower.endswith('.xls'):
            extracted = self._parse_excel(file_content)
        elif filename_lower.endswith('.txt'):
            text = file_content.decode('utf-8', errors='ignore')
            extracted = self._extract_from_text(text)
        else:
            extracted = {}
        
        # Find missing required fields
        missing = self._get_missing_fields(extracted)
        
        return extracted, missing
    
    def parse_document_multi(self, file_content: bytes, filename: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """
        Parse a document and extract data for multiple people (e.g., couples).
        
        Returns:
            Tuple of (list of person dicts, shared data dict like address/portfolio)
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.docx'):
            text = self._get_docx_text(file_content)
        elif filename_lower.endswith('.txt'):
            text = file_content.decode('utf-8', errors='ignore')
        else:
            # For other formats, fall back to single extraction
            extracted, _ = self.parse_document(file_content, filename)
            return [extracted], {}
        
        return self._extract_multiple_people(text)
    
    def _get_docx_text(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        if not self._docx_available:
            return ""
        
        try:
            from docx import Document
            doc = Document(BytesIO(file_content))
            
            # Extract all text from paragraphs
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Also check tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    full_text += "\n" + row_text
            
            return full_text
        except Exception as e:
            return ""
    
    def _extract_multiple_people(self, text: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """
        Extract data for multiple people from text.
        Detects couples/partners format like "Johns & Willy Johns"
        """
        people = []
        shared = {}
        
        # Extract shared data (address, portfolio, etc.)
        # These apply to the household
        
        # Address
        address_match = re.search(
            r'Address[:\s]*([^,\n]+(?:,\s*[^,\n]+)*,\s*[A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2})',
            text, re.IGNORECASE
        )
        if address_match:
            addr_parts = address_match.group(1).split(',')
            if len(addr_parts) >= 2:
                shared["address_line1"] = addr_parts[0].strip()
                # Find city (usually second-to-last before postcode)
                if len(addr_parts) >= 3:
                    shared["city"] = addr_parts[-2].strip()
                # Postcode is in the last part
                postcode_match = re.search(r'([A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2})', addr_parts[-1].upper())
                if postcode_match:
                    shared["postcode"] = postcode_match.group(1)
        
        # Net Worth / Portfolio
        networth_match = re.search(r'Net\s*Worth[:\s]*£([\d,]+)', text, re.IGNORECASE)
        if networth_match:
            shared["portfolio"] = self._clean_number(networth_match.group(1))
        
        # Detect couple format: "CLIENT X: NAME1 & NAME2 SURNAME"
        couple_match = re.search(
            r'CLIENT\s*\d*[:\s]+([A-Z][A-Za-z]+)\s*(?:&|AND)\s*([A-Z][A-Za-z]+)\s+([A-Z][A-Za-z]+)',
            text, re.IGNORECASE
        )
        
        if couple_match:
            name1 = couple_match.group(1).strip().title()
            name2 = couple_match.group(2).strip().title()
            surname = couple_match.group(3).strip().title()
            
            # Extract Person 1
            person1 = {"first_name": name1, "last_name": surname}
            
            # Find Person 1's DOB - look for "Name1 Name1" or just "Name1" followed by DOB
            dob1_match = re.search(
                rf'{name1}(?:\s+{surname})?\s*\n\s*DOB[:\s]*(\d{{1,2}}[\/\-\.]\d{{1,2}}[\/\-\.]\d{{2,4}})',
                text, re.IGNORECASE
            )
            if dob1_match:
                person1["date_of_birth"] = self._parse_date(dob1_match.group(1))
            
            # Find Person 1's contact
            contact1_match = re.search(
                rf'{name1}[:\s]*(07\d{{3}}\s*\d{{3}}\s*\d{{3,4}})\s*\|?\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{{2,}})?',
                text, re.IGNORECASE
            )
            if contact1_match:
                person1["phone"] = contact1_match.group(1)
                if contact1_match.group(2):
                    person1["email"] = contact1_match.group(2)
            
            # Find Person 1's occupation and income
            occ1_match = re.search(
                rf'{name1}[:\s]+([A-Za-z][A-Za-z\s,]+?)\s*[–-]\s*£([\d,]+)',
                text, re.IGNORECASE
            )
            if occ1_match:
                person1["occupation"] = occ1_match.group(1).strip().rstrip(',')
                person1["income"] = self._clean_number(occ1_match.group(2))
            
            people.append(person1)
            
            # Extract Person 2
            person2 = {"first_name": name2, "last_name": surname}
            
            # Find Person 2's DOB
            dob2_match = re.search(
                rf'{name2}(?:\s+{surname})?\s*\n\s*DOB[:\s]*(\d{{1,2}}[\/\-\.]\d{{1,2}}[\/\-\.]\d{{2,4}})',
                text, re.IGNORECASE
            )
            if dob2_match:
                person2["date_of_birth"] = self._parse_date(dob2_match.group(1))
            
            # Find Person 2's contact
            contact2_match = re.search(
                rf'{name2}[:\s]*(07\d{{3}}\s*\d{{3}}\s*\d{{3,4}})',
                text, re.IGNORECASE
            )
            if contact2_match:
                person2["phone"] = contact2_match.group(1)
            
            # Find Person 2's occupation and income
            occ2_match = re.search(
                rf'{name2}[:\s]+([A-Za-z][A-Za-z\s,]+?)\s*[–-]\s*£([\d,]+)',
                text, re.IGNORECASE
            )
            if occ2_match:
                person2["occupation"] = occ2_match.group(1).strip().rstrip(',')
                person2["income"] = self._clean_number(occ2_match.group(2))
            
            people.append(person2)
        
        else:
            # Single person - use standard extraction
            extracted = self._extract_from_text(text)
            people.append(extracted)
        
        return people, shared

    def _parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Parse Word document"""
        text = self._get_docx_text(file_content)
        if not text:
            return {"_error": "python-docx not installed or error reading file"}
        return self._extract_from_text(text)
    
    def _parse_excel(self, file_content: bytes) -> Dict[str, Any]:
        """Parse Excel document"""
        if not self._openpyxl_available:
            return {"_error": "openpyxl not installed"}
        
        try:
            from openpyxl import load_workbook
            wb = load_workbook(BytesIO(file_content), data_only=True)
            
            extracted = {}
            
            # Check first sheet
            ws = wb.active
            
            # Build text from all cells and also check for key-value pairs
            full_text = ""
            for row in ws.iter_rows():
                row_values = []
                prev_cell = None
                
                for cell in row:
                    if cell.value is not None:
                        cell_str = str(cell.value).strip()
                        row_values.append(cell_str)
                        
                        # Check if previous cell was a label
                        if prev_cell and self._is_label(prev_cell):
                            field = self._label_to_field(prev_cell)
                            if field:
                                extracted[field] = cell_str
                        
                        prev_cell = cell_str
                    else:
                        prev_cell = None
                
                full_text += " | ".join(row_values) + "\n"
            
            # Also do pattern matching on full text
            text_extracted = self._extract_from_text(full_text)
            
            # Merge (direct extraction takes precedence)
            for key, value in text_extracted.items():
                if key not in extracted:
                    extracted[key] = value
            
            return extracted
            
        except Exception as e:
            return {"_error": f"Error parsing Excel document: {str(e)}"}
    
    def _parse_text(self, text: str) -> Dict[str, Any]:
        """Parse plain text"""
        return self._extract_from_text(text)
    
    def _extract_from_text(self, text: str) -> Dict[str, Any]:
        """Extract client data from text using pattern matching"""
        extracted = {}
        text_lower = text.lower()
        
        # ===== SPECIAL: Extract name from "CLIENT X: NAME & PARTNER NAME" format =====
        client_header_match = re.search(
            r'CLIENT\s*\d*[:\s]+([A-Z][A-Za-z]+)(?:\s*(?:&|AND)\s*([A-Z][A-Za-z]+))?\s+([A-Z][A-Za-z]+)',
            text, re.IGNORECASE
        )
        if client_header_match:
            # Format: "CLIENT 13: JOHNS & WILLY JOHNS" -> first_name=Johns, last_name=Johns
            first = client_header_match.group(1)
            last = client_header_match.group(3)
            if first and last:
                extracted["first_name"] = first.strip().title()
                extracted["last_name"] = last.strip().title()
        
        # ===== SPECIAL: Extract name after "PERSONAL DETAILS" section =====
        if "first_name" not in extracted:
            personal_match = re.search(
                r'PERSONAL\s+DETAILS\s*\n+([A-Z][a-z]+)\s+([A-Z][a-z]+)',
                text, re.IGNORECASE
            )
            if personal_match:
                extracted["first_name"] = personal_match.group(1).strip().title()
                extracted["last_name"] = personal_match.group(2).strip().title()
        
        # ===== SPECIAL: Look for "Name Name" followed by "DOB:" pattern =====
        if "first_name" not in extracted:
            name_dob_match = re.search(
                r'([A-Z][a-z]+)\s+([A-Z][a-z]+)\s*\n\s*(?:DOB|Date\s*of\s*Birth)',
                text, re.IGNORECASE
            )
            if name_dob_match:
                extracted["first_name"] = name_dob_match.group(1).strip().title()
                extracted["last_name"] = name_dob_match.group(2).strip().title()
        
        # ===== SPECIAL: Extract address from format "Address: Street, City, POSTCODE" =====
        address_match = re.search(
            r'(?:Address|Residence)[:\s]*([^,\n]+),\s*([^,\n]+),\s*([^,\n]+),\s*([A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2})',
            text, re.IGNORECASE
        )
        if address_match:
            extracted["address_line1"] = address_match.group(1).strip()
            # City might be in group 2 or 3
            city_candidate = address_match.group(3).strip()
            if len(city_candidate) < 3:
                city_candidate = address_match.group(2).strip()
            extracted["city"] = city_candidate
            extracted["postcode"] = address_match.group(4).strip().upper()
        
        # Try each standard pattern
        for field, pattern in self.FIELD_PATTERNS.items():
            if field in extracted:  # Skip if already found
                continue
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = match.group(1).strip()
                
                # Clean up the value based on field type
                if field == "income" or field == "portfolio":
                    value = self._clean_number(value)
                elif field == "date_of_birth":
                    value = self._parse_date(value)
                elif field == "full_name" and "first_name" not in extracted:
                    # Split full name
                    parts = value.split()
                    if len(parts) >= 2:
                        # Check if first part is a title
                        if parts[0] in self.VALID_TITLES:
                            extracted["title"] = parts[0]
                            extracted["first_name"] = parts[1] if len(parts) > 1 else ""
                            extracted["last_name"] = " ".join(parts[2:]) if len(parts) > 2 else ""
                        else:
                            extracted["first_name"] = parts[0]
                            extracted["last_name"] = " ".join(parts[1:])
                    continue
                elif field == "marital":
                    value = self._normalize_marital_status(value)
                elif field == "title":
                    value = self._normalize_title(value)
                elif field == "postcode":
                    value = value.upper()
                elif field == "address":
                    extracted["address_line1"] = value
                    continue
                
                if value:
                    extracted[field] = value
        
        # Try to find email anywhere in text
        if "email" not in extracted:
            email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
            if email_match:
                extracted["email"] = email_match.group(0)
        
        # Try to find UK postcode anywhere
        if "postcode" not in extracted:
            postcode_match = re.search(self.UK_POSTCODE, text.upper())
            if postcode_match:
                extracted["postcode"] = postcode_match.group(0)
        
        # Try to find phone number (UK mobile format)
        if "phone" not in extracted:
            phone_match = re.search(r'(07\d{3}\s*\d{3}\s*\d{3,4})', text)
            if phone_match:
                extracted["phone"] = phone_match.group(1)
        
        # Try to extract occupation from "Name: Job Title – £amount" format
        if "occupation" not in extracted:
            # Pattern: "Johns: Regional Radio Presenter – £68,000"
            occupation_match = re.search(
                r'[A-Za-z]+:\s+([A-Za-z][A-Za-z\s,]+?)\s*[–-]\s*£[\d,]+',
                text
            )
            if occupation_match:
                extracted["occupation"] = occupation_match.group(1).strip().rstrip(',')
        
        # Try to extract income from "– £68,000" format
        if "income" not in extracted:
            income_match = re.search(r'[–-]\s*£([\d,]+)', text)
            if income_match:
                extracted["income"] = self._clean_number(income_match.group(1))
        
        # Try to extract net worth / portfolio
        if "portfolio" not in extracted:
            networth_match = re.search(r'Net\s*Worth[:\s]*£([\d,]+)', text, re.IGNORECASE)
            if networth_match:
                extracted["portfolio"] = self._clean_number(networth_match.group(1))
        
        # Try to extract city from address line
        if "city" not in extracted and "postcode" in extracted:
            # Look for city name before postcode
            city_match = re.search(
                r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*' + extracted["postcode"],
                text, re.IGNORECASE
            )
            if city_match:
                extracted["city"] = city_match.group(1).strip()
        
        return extracted
    
    def _is_label(self, text: str) -> bool:
        """Check if text looks like a field label"""
        labels = [
            "name", "first", "last", "surname", "email", "phone", "tel",
            "address", "city", "postcode", "dob", "birth", "occupation",
            "employer", "income", "salary", "title", "marital"
        ]
        text_lower = text.lower()
        return any(label in text_lower for label in labels)
    
    def _label_to_field(self, label: str) -> Optional[str]:
        """Convert a label to a field name"""
        label_lower = label.lower()
        mapping = {
            "first name": "first_name",
            "last name": "last_name",
            "surname": "last_name",
            "forename": "first_name",
            "email": "email",
            "e-mail": "email",
            "phone": "phone",
            "telephone": "phone",
            "mobile": "phone",
            "address": "address_line1",
            "city": "city",
            "town": "city",
            "postcode": "postcode",
            "post code": "postcode",
            "date of birth": "date_of_birth",
            "dob": "date_of_birth",
            "occupation": "occupation",
            "job": "occupation",
            "employer": "employer",
            "company": "employer",
            "income": "income",
            "salary": "income",
            "title": "title",
            "marital status": "marital_status",
        }
        
        for key, field in mapping.items():
            if key in label_lower:
                return field
        return None
    
    def _clean_number(self, value: str) -> Optional[float]:
        """Clean and convert a number string"""
        try:
            # Remove currency symbols, commas, spaces
            cleaned = re.sub(r'[£$,\s]', '', value)
            return float(cleaned)
        except:
            return None
    
    def _parse_date(self, value: str) -> Optional[str]:
        """Parse various date formats to ISO format"""
        formats = [
            "%d/%m/%Y", "%d-%m-%Y", "%d.%m.%Y",
            "%d/%m/%y", "%d-%m-%y", "%d.%m.%y",
            "%Y-%m-%d", "%Y/%m/%d",
        ]
        
        for fmt in formats:
            try:
                parsed = datetime.strptime(value, fmt)
                return parsed.date().isoformat()
            except:
                continue
        return value
    
    def _normalize_marital_status(self, value: str) -> str:
        """Normalize marital status value"""
        value_lower = value.lower()
        if "single" in value_lower:
            return "single"
        elif "married" in value_lower:
            return "married"
        elif "divorced" in value_lower:
            return "divorced"
        elif "widow" in value_lower:
            return "widowed"
        elif "civil" in value_lower:
            return "civil_partnership"
        return "single"
    
    def _normalize_title(self, value: str) -> str:
        """Normalize title"""
        value_cap = value.strip().capitalize()
        if value_cap in self.VALID_TITLES:
            return value_cap
        # Common variations
        if value_cap in ["Mister", "Master"]:
            return "Mr"
        if value_cap in ["Missus", "Misses"]:
            return "Mrs"
        if value_cap == "Doctor":
            return "Dr"
        return "Mr"  # Default
    
    def _get_missing_fields(self, extracted: Dict[str, Any]) -> List[str]:
        """Get list of required fields that are missing"""
        missing = []
        
        # Map extracted fields to required fields
        field_mapping = {
            "id": "id",
            "title": "title",
            "first_name": "first_name",
            "last_name": "last_name",
            "date_of_birth": "date_of_birth",
            "email": "email",
            "phone": "phone",
            "address_line1": "address_line1",
            "city": "city",
            "postcode": "postcode",
        }
        
        for required in self.REQUIRED_FIELDS:
            if required == "id":
                # ID is always auto-generated
                continue
            if required not in extracted or not extracted.get(required):
                missing.append(required)
        
        return missing
    
    def generate_client_id(self, existing_ids: List[str]) -> str:
        """Generate a unique client ID"""
        max_num = 0
        for cid in existing_ids:
            match = re.search(r'client_(\d+)', cid)
            if match:
                num = int(match.group(1))
                max_num = max(max_num, num)
        
        return f"client_{max_num + 1:03d}"


# Singleton instance
document_parser = DocumentParser()
